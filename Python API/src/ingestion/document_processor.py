"""
Document ingestion module for processing and storing documents.
Handles PDF, text, markdown, and DOCX files.
"""

from typing import List, Dict, Any
from pathlib import Path
import hashlib
from loguru import logger
from PyPDF2 import PdfReader
from docx import Document
import markdown
from bs4 import BeautifulSoup
import io


class DocumentProcessor:
    """Processes various document formats into text chunks."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize document processor.
        
        Args:
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between consecutive chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def process_file(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Process a file and return text chunks with metadata.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of document chunks with metadata
        """
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            text = self._extract_pdf(file_path)
        elif suffix == '.txt':
            text = self._extract_text(file_path)
        elif suffix == '.md':
            text = self._extract_markdown(file_path)
        elif suffix == '.docx':
            text = self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
        
        chunks = self._create_chunks(text)
        
        return [
            {
                "text": chunk,
                "metadata": {
                    "source": str(file_path),
                    "filename": file_path.name,
                    "chunk_index": idx,
                    "file_type": suffix,
                    "doc_id": self._generate_doc_id(file_path)
                }
            }
            for idx, chunk in enumerate(chunks)
        ]
    
    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        logger.info(f"Extracting text from PDF: {file_path}")
        reader = PdfReader(str(file_path))
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    
    def _extract_text(self, file_path: Path) -> str:
        """Extract text from plain text file."""
        logger.info(f"Reading text file: {file_path}")
        return file_path.read_text(encoding='utf-8')
    
    def _extract_markdown(self, file_path: Path) -> str:
        """Extract text from markdown file."""
        logger.info(f"Processing markdown file: {file_path}")
        md_content = file_path.read_text(encoding='utf-8')
        html = markdown.markdown(md_content)
        soup = BeautifulSoup(html, 'html.parser')
        return soup.get_text()
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        logger.info(f"Extracting text from DOCX: {file_path}")
        doc = Document(str(file_path))
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    def _create_chunks(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks.
        
        Args:
            text: Input text to chunk
            
        Returns:
            List of text chunks
        """
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = start + self.chunk_size
            chunk = text[start:end]
            
            # Try to end chunk at sentence boundary
            if end < text_length:
                last_period = chunk.rfind('.')
                last_newline = chunk.rfind('\n')
                boundary = max(last_period, last_newline)
                
                if boundary > self.chunk_size // 2:
                    chunk = chunk[:boundary + 1]
                    end = start + boundary + 1
            
            chunks.append(chunk.strip())
            start = end - self.chunk_overlap
        
        return [chunk for chunk in chunks if chunk]
    
    def _generate_doc_id(self, file_path: Path) -> str:
        """Generate unique document ID from file path."""
        return hashlib.md5(str(file_path).encode()).hexdigest()
    
    def process_upload(self, filename: str, file_content: bytes) -> List[Dict[str, Any]]:
        """
        Process an uploaded file from bytes without saving to disk.
        
        Args:
            filename: Original filename
            file_content: File content as bytes
            
        Returns:
            List of document chunks with metadata
        """
        file_extension = Path(filename).suffix.lower()
        
        if file_extension == '.pdf':
            text = self._extract_pdf_from_bytes(file_content)
        elif file_extension == '.txt':
            text = self._extract_text_from_bytes(file_content)
        elif file_extension == '.md':
            text = self._extract_markdown_from_bytes(file_content)
        elif file_extension == '.docx':
            text = self._extract_docx_from_bytes(file_content)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        chunks = self._create_chunks(text)
        
        # Generate doc_id from file content hash (not filename) for true deduplication
        doc_id = hashlib.md5(file_content).hexdigest()
        
        return [
            {
                "text": chunk,
                "metadata": {
                    "source": filename,
                    "filename": filename,
                    "chunk_index": idx,
                    "file_type": file_extension,
                    "doc_id": doc_id
                }
            }
            for idx, chunk in enumerate(chunks)
        ]
    
    def _extract_pdf_from_bytes(self, file_content: bytes) -> str:
        """Extract text from PDF bytes."""
        logger.info("Extracting text from PDF")
        reader = PdfReader(io.BytesIO(file_content))
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    
    def _extract_text_from_bytes(self, file_content: bytes) -> str:
        """Extract text from plain text bytes."""
        logger.info("Reading text file")
        return file_content.decode('utf-8')
    
    def _extract_markdown_from_bytes(self, file_content: bytes) -> str:
        """Extract text from markdown bytes."""
        logger.info("Processing markdown file")
        md_content = file_content.decode('utf-8')
        html = markdown.markdown(md_content)
        soup = BeautifulSoup(html, 'html.parser')
        return soup.get_text()
    
    def _extract_docx_from_bytes(self, file_content: bytes) -> str:
        """Extract text from DOCX bytes."""
        logger.info("Extracting text from DOCX")
        doc = Document(io.BytesIO(file_content))
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
